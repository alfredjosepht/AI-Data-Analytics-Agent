from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import os

from backend.agents.query_agent.query_agent import QueryAgent, sanitize_json_values
from backend.database.sqlite_manager import sqlite_manager

router = APIRouter(
    tags=["Chat"]
)


class ChatRequest(BaseModel):
    question: str


@router.post("/")
def chat(request: ChatRequest):
    return QueryAgent.execute_question(request.question)


@router.get("/workspaces/{workspace_id}/chat_history")
def get_chat_history(workspace_id: int):
    import json
    history = sqlite_manager.get_chat_history(workspace_id)
    parsed_history = []
    for chat in history:
        res_data = None
        if chat.get("result_json"):
            try:
                res_data = json.loads(chat["result_json"])
            except Exception:
                res_data = []
        
        chart_data = None
        if chat.get("chart_json"):
            try:
                chart_data = json.loads(chat["chart_json"])
            except Exception:
                chart_data = None

        parsed_history.append({
            "id": chat.get("id"),
            "workspace_id": chat.get("workspace_id"),
            "question": chat.get("question"),
            "answer": chat.get("answer"),
            "sql_query": chat.get("sql_query"),
            "result": res_data,
            "chart": chart_data,
            "created_at": chat.get("created_at")
        })
    return sanitize_json_values({
        "status": "success",
        "chat_history": parsed_history
    })


@router.get("/workspaces/{workspace_id}/chat_messages")
def get_chat_messages(workspace_id: int):
    import json
    messages = sqlite_manager.get_chat_messages(workspace_id)
    parsed_messages = []
    for msg in messages:
        res_data = None
        if msg.get("result_json"):
            try:
                res_data = json.loads(msg["result_json"])
            except Exception:
                res_data = []
        
        chart_data = None
        if msg.get("chart_json"):
            try:
                chart_data = json.loads(msg["chart_json"])
            except Exception:
                chart_data = None

        parsed_messages.append({
            "id": msg.get("id"),
            "workspace_id": msg.get("workspace_id"),
            "role": msg.get("role"),
            "message": msg.get("message"),
            "sql_query": msg.get("sql_query"),
            "result": res_data,
            "chart": chart_data,
            "created_at": msg.get("created_at")
        })
    return sanitize_json_values({
        "status": "success",
        "chat_messages": parsed_messages
    })


@router.delete("/workspaces/{workspace_id}/chat_history")
def clear_chat_history(workspace_id: int):
    sqlite_manager.clear_workspace_chat(workspace_id)
    return {
        "status": "success",
        "message": f"Chat history cleared for workspace {workspace_id}"
    }


@router.delete("/chat_history/{chat_id}")
def delete_chat_message(chat_id: int):
    sqlite_manager.delete_chat_message(chat_id)
    return {
        "status": "success",
        "message": f"Message {chat_id} deleted."
    }


@router.get("/workspaces/{workspace_id}/chat_history/export")
def export_chat_history(workspace_id: int, format: str = "txt"):
    workspace = sqlite_manager.get_workspace(workspace_id)
    if not workspace:
        raise HTTPException(status_code=404, detail="Workspace not found")
        
    history = sqlite_manager.get_chat_history(workspace_id)
    out_dir = "exports"
    os.makedirs(out_dir, exist_ok=True)
    
    if format == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        
        out_path = os.path.join(out_dir, f"chat_history_{workspace_id}.pdf")
        doc = SimpleDocTemplate(out_path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1e293b'),
            spaceAfter=15
        )
        msg_style_user = ParagraphStyle(
            'UserMsg',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#2563eb'),
            spaceAfter=6
        )
        msg_style_agent = ParagraphStyle(
            'AgentMsg',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceAfter=12
        )
        
        elements = []
        elements.append(Paragraph(f"Chat History - Workspace: {workspace.get('name').upper()}", title_style))
        elements.append(Spacer(1, 10))
        
        for idx, chat in enumerate(history):
            q_time = chat.get('created_at', '')
            elements.append(Paragraph(f"<b>User ({q_time}):</b> {chat.get('question')}", msg_style_user))
            elements.append(Paragraph(f"<b>AI Insight:</b> {chat.get('answer')}", msg_style_agent))
            if chat.get('sql_query'):
                elements.append(Paragraph(f"<b>SQL Query:</b> <code>{chat.get('sql_query')}</code>", msg_style_agent))
            elements.append(Spacer(1, 10))
            
        doc.build(elements)
        return FileResponse(out_path, media_type="application/pdf", filename=f"chat_history_{workspace.get('name')}.pdf")
        
    elif format == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        out_path = os.path.join(out_dir, f"chat_history_{workspace_id}.docx")
        doc = Document()
        
        # Heading
        heading = doc.add_heading(level=1)
        run = heading.add_run(f"Chat History - Workspace: {workspace.get('name').upper()}")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(30, 41, 59)
        
        doc.add_paragraph(f"Source Dataset: {workspace.get('file_name')}")
        doc.add_paragraph()
        
        for idx, chat in enumerate(history):
            q_time = chat.get('created_at', '')
            p1 = doc.add_paragraph()
            run_user = p1.add_run(f"User ({q_time}): {chat.get('question')}")
            run_user.bold = True
            run_user.font.color.rgb = RGBColor(37, 99, 235)
            
            p2 = doc.add_paragraph()
            run_agent = p2.add_run(f"AI Insight: {chat.get('answer')}")
            run_agent.font.color.rgb = RGBColor(51, 65, 85)
            
            if chat.get('sql_query'):
                p3 = doc.add_paragraph()
                run_sql = p3.add_run(f"SQL: {chat.get('sql_query')}")
                run_sql.font.name = 'Courier New'
                run_sql.font.size = Pt(9)
            
            doc.add_paragraph()
            
        doc.save(out_path)
        return FileResponse(out_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"chat_history_{workspace.get('name')}.docx")
        
    else: # TXT format
        out_path = os.path.join(out_dir, f"chat_history_{workspace_id}.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(f"CHAT HISTORY - WORKSPACE: {workspace.get('name').upper()}\n")
            f.write(f"Source Dataset: {workspace.get('file_name')}\n")
            f.write("=" * 60 + "\n\n")
            for idx, chat in enumerate(history):
                q_time = chat.get('created_at', '')
                f.write(f"[{idx+1}] User ({q_time}): {chat.get('question')}\n")
                f.write(f"AI: {chat.get('answer')}\n")
                if chat.get('sql_query'):
                    f.write(f"SQL: {chat.get('sql_query')}\n")
                f.write("-" * 40 + "\n\n")
                
        return FileResponse(out_path, media_type="text/plain", filename=f"chat_history_{workspace.get('name')}.txt")
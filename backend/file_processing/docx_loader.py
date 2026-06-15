from docx import Document
import pandas as pd


class DOCXLoader:

    @staticmethod
    def load(file_path: str):

        try:

            document = Document(file_path)

            paragraphs = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]

            tables = []

            for table in document.tables:
                rows = [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ]

                if len(rows) > 1:
                    df = pd.DataFrame(
                        rows[1:],
                        columns=rows[0]
                    )
                    tables.append(df)

            return {
                "text": "\n".join(paragraphs),
                "tables": tables
            }

        except Exception as e:

            raise Exception(
                f"DOCX loading failed: {str(e)}"
            )
